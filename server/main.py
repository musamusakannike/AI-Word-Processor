from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from google import genai
from google.genai import types
import os
import uuid
from pathlib import Path
import traceback
from datetime import datetime
from dotenv import load_dotenv
import asyncio
from contextlib import asynccontextmanager
from html.parser import HTMLParser
from docx import Document
import io
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.enums import TA_LEFT, TA_CENTER

# Load environment variables from .env file
load_dotenv()

# Create directories for storing generated files
GENERATED_DIR = Path("generated_files")
GENERATED_DIR.mkdir(exist_ok=True)

# File upload constants
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB
ALLOWED_FILE_TYPES = {"application/pdf", "text/plain", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
ALLOWED_EXTENSIONS = {".pdf", ".txt", ".docx"}

async def periodic_cleanup():
    print("Cleanup task started")
    while True:
        await asyncio.sleep(60)  # Check every minute, wait before first check
        try:
            current_time = datetime.now().timestamp()
            max_age_seconds = 600  # 10 minutes (10 * 60)
            
            for file_path in GENERATED_DIR.glob("*"):
                if not file_path.suffix.lower() in [".docx", ".pdf"]:
                    continue
                try:
                    if not file_path.exists():
                        continue
                        
                    stats = file_path.stat()
                    file_age = current_time - stats.st_mtime
                    
                    if file_age > max_age_seconds:
                        file_path.unlink()
                        print(f"Deleted old file: {file_path.name}")
                except OSError as e:
                    print(f"Error accessing/deleting file {file_path}: {e}")
                except Exception as e:
                    print(f"Unexpected error with file {file_path}: {e}")
                    
        except Exception as e:
            print(f"Error in cleanup loop: {e}")

@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup: Start the background task
    cleanup_task = asyncio.create_task(periodic_cleanup())
    yield
    # Shutdown: Cancel the background task
    cleanup_task.cancel()
    try:
        await cleanup_task
    except asyncio.CancelledError:
        pass

app = FastAPI(title="AI Word Processor API", lifespan=lifespan)

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)




# Mount the generated files directory for serving files
app.mount("/files", StaticFiles(directory=str(GENERATED_DIR)), name="files")

# Configure Gemini API
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise ValueError("GEMINI_API_KEY environment variable is not set")

client = genai.Client(api_key=GEMINI_API_KEY)
MODEL_NAME = "gemini-3-flash-preview"

DOC_GEN_SYS_INSTRUCT = """You are an expert document automation engineer specializing in `python-docx`.
Your goal is to generate Python code that creates highly professional, visually appealing, and comprehensive DOCX documents.

RULES:
1. **Output ONLY executable Python code**. No markdown backticks, no explanations, no comments.
2. **Professional Styling**:
   - Use a clean, modern font (e.g., Arial, Calibri, or Open Sans) if possible, or stick to standard professional fonts.
   - Use appropriate font sizes: 14pt-16pt for headings, 11pt-12pt for body text.
   - Add proper spacing between paragraphs (e.g., `paragraph_format.space_after = Pt(12)`).
   - Use bold and italics for emphasis where appropriate.
3. **Structure & Length**:
   - The document must be **EXTENSIVE** and **DETAILED**. Expand significantly on the user's prompt.
   - Use a clear hierarchy: Title -> Heading 1 -> Heading 2 -> Body Text.
   - Include clear sections (e.g., Introduction, Detailed Analysis, Key Findings, Conclusion).
   - Incorporate lists (bullet points and numbered lists) to break up text.
   - Use tables for structured data if relevant.
4. **Execution**:
   - The code must be complete and error-free.
   - Use the variable `output_path` for saving the file.
   - Imports must include all necessary classes from `docx`, `docx.shared`, `docx.enum.text`.

Example of expected code structure:
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Professional Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Section
doc.add_heading('1. Introduction', level=1)
p = doc.add_paragraph('This is a detailed introduction...')
p.paragraph_format.space_after = Pt(12)

# Save
doc.save(output_path)
"""

EDIT_SYS_INSTRUCT = """You are an expert Word document editor.
You will be given an HTML document and an instruction.

RULES:
1. Output ONLY valid HTML.
2. Do not wrap the output in markdown backticks.
3. Preserve the document structure as much as possible.
4. Apply the instruction precisely.
"""


def _strip_code_fences(text: str) -> str:
    cleaned = text.strip()
    if cleaned.startswith("```html"):
        cleaned = cleaned.replace("```html", "", 1)
    elif cleaned.startswith("```"):
        cleaned = cleaned.replace("```", "", 1)
    if cleaned.endswith("```"):
        cleaned = cleaned[:-3]
    return cleaned.strip()


class _HtmlToDocxParser(HTMLParser):
    def __init__(self, doc: Document):
        super().__init__(convert_charrefs=True)
        self._doc = doc
        self._block_tag: str | None = None
        self._heading_level: int | None = None
        self._list_mode: str | None = None
        self._buffer: list[str] = []

    def handle_starttag(self, tag: str, attrs):
        if tag in {"p", "h1", "h2", "h3", "li", "blockquote"}:
            self._flush_block()
            self._block_tag = tag
            if tag == "h1":
                self._heading_level = 1
            elif tag == "h2":
                self._heading_level = 2
            elif tag == "h3":
                self._heading_level = 3
            else:
                self._heading_level = None
        elif tag == "ul":
            self._list_mode = "bullet"
        elif tag == "ol":
            self._list_mode = "number"
        elif tag == "br":
            self._buffer.append("\n")

    def handle_endtag(self, tag: str):
        if tag in {"p", "h1", "h2", "h3", "li", "blockquote"}:
            self._flush_block()
            self._block_tag = None
            self._heading_level = None
        elif tag in {"ul", "ol"}:
            self._list_mode = None

    def handle_data(self, data: str):
        if not data:
            return
        self._buffer.append(data)

    def _flush_block(self):
        text = "".join(self._buffer).strip()
        self._buffer = []
        if not text:
            return

        if self._heading_level is not None:
            self._doc.add_heading(text, level=self._heading_level)
            return

        if self._block_tag == "li":
            style = "List Bullet" if self._list_mode == "bullet" else "List Number"
            self._doc.add_paragraph(text, style=style)
            return

        if self._block_tag == "blockquote":
            self._doc.add_paragraph(text, style="Intense Quote")
            return

        self._doc.add_paragraph(text)


def html_to_docx_bytes(html: str) -> bytes:
    doc = Document()
    parser = _HtmlToDocxParser(doc)
    parser.feed(html)
    parser.close()

    tmp_id = str(uuid.uuid4())
    tmp_path = GENERATED_DIR / f"tmp_{tmp_id}.docx"
    doc.save(str(tmp_path))
    data = tmp_path.read_bytes()
    tmp_path.unlink(missing_ok=True)
    return data


class _HtmlToPdfParser(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self._story: list = []
        self._styles = getSampleStyleSheet()
        self._block_tag: str | None = None
        self._heading_level: int | None = None
        self._list_mode: str | None = None
        self._buffer: list[str] = []
        self._list_items: list = []

    def handle_starttag(self, tag: str, attrs):
        if tag in {"p", "h1", "h2", "h3", "li", "blockquote"}:
            self._flush_block()
            self._block_tag = tag
            if tag == "h1":
                self._heading_level = 1
            elif tag == "h2":
                self._heading_level = 2
            elif tag == "h3":
                self._heading_level = 3
            else:
                self._heading_level = None
        elif tag == "ul":
            self._list_mode = "bullet"
        elif tag == "ol":
            self._list_mode = "number"
        elif tag == "br":
            self._buffer.append("<br/>")

    def handle_endtag(self, tag: str):
        if tag in {"p", "h1", "h2", "h3", "li", "blockquote"}:
            self._flush_block()
            self._block_tag = None
            self._heading_level = None
        elif tag in {"ul", "ol"}:
            if self._list_items:
                list_flowable = ListFlowable(
                    self._list_items,
                    bulletType='bullet' if self._list_mode == "bullet" else '1',
                    leftIndent=20,
                    bulletOffsetY=-2
                )
                self._story.append(list_flowable)
                self._story.append(Spacer(1, 0.1 * inch))
                self._list_items = []
            self._list_mode = None

    def handle_data(self, data: str):
        if not data:
            return
        self._buffer.append(data)

    def _flush_block(self):
        text = "".join(self._buffer).strip()
        self._buffer = []
        if not text:
            return

        if self._heading_level == 1:
            style = self._styles['Heading1']
            self._story.append(Paragraph(text, style))
            self._story.append(Spacer(1, 0.2 * inch))
        elif self._heading_level == 2:
            style = self._styles['Heading2']
            self._story.append(Paragraph(text, style))
            self._story.append(Spacer(1, 0.15 * inch))
        elif self._heading_level == 3:
            style = self._styles['Heading3']
            self._story.append(Paragraph(text, style))
            self._story.append(Spacer(1, 0.1 * inch))
        elif self._block_tag == "li":
            self._list_items.append(Paragraph(text, self._styles['BodyText']))
        elif self._block_tag == "blockquote":
            quote_style = ParagraphStyle(
                'Quote',
                parent=self._styles['BodyText'],
                leftIndent=20,
                rightIndent=20,
                textColor='#666666',
                fontName='Helvetica-Oblique'
            )
            self._story.append(Paragraph(text, quote_style))
            self._story.append(Spacer(1, 0.1 * inch))
        else:
            self._story.append(Paragraph(text, self._styles['BodyText']))
            self._story.append(Spacer(1, 0.1 * inch))

    def get_story(self):
        self._flush_block()
        return self._story


def html_to_pdf_bytes(html: str) -> bytes:
    tmp_id = str(uuid.uuid4())
    tmp_path = GENERATED_DIR / f"tmp_{tmp_id}.pdf"
    
    doc = SimpleDocTemplate(
        str(tmp_path),
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    parser = _HtmlToPdfParser()
    parser.feed(html)
    parser.close()
    
    story = parser.get_story()
    if not story:
        story = [Paragraph("Empty document", getSampleStyleSheet()['BodyText'])]
    
    doc.build(story)
    
    data = tmp_path.read_bytes()
    tmp_path.unlink(missing_ok=True)
    return data


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    return "\n".join(text_parts)


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from a TXT file."""
    try:
        return file_bytes.decode('utf-8')
    except UnicodeDecodeError:
        # Try with latin-1 as fallback
        return file_bytes.decode('latin-1')


def validate_upload_file(file: UploadFile) -> None:
    """Validate uploaded file type and size."""
    if not file:
        return
    
    # Check file extension
    file_ext = Path(file.filename or "").suffix.lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
        )
    
    # Note: Size validation will be done after reading the file


class DocumentRequest(BaseModel):
    prompt: str


class DocumentResponse(BaseModel):
    success: bool
    message: str
    download_url: str | None = None
    filename: str | None = None
    generated_code: str | None = None
    error: str | None = None


class ExportRequest(BaseModel):
    html: str
    filename: str | None = None


class AiEditRequest(BaseModel):
    html: str
    instruction: str


class AiEditResponse(BaseModel):
    success: bool
    message: str
    updated_html: str | None = None
    error: str | None = None


class Template(BaseModel):
    id: str
    title: str
    description: str
    category: str
    prompt: str
    icon: str


TEMPLATES = [
    Template(
        id="modern-resume",
        title="Modern Resume",
        description="Professional resume with clean design and modern formatting",
        category="Professional",
        prompt="""Create a Modern Resume using the following specifications:

FORMATTING REQUIREMENTS:
- Use Calibri font, 11pt for body text, 14pt for name/header
- Professional blue header (#2563EB) with white text for the name section
- Clear section headings in bold, 12pt
- Consistent spacing: 12pt after paragraphs, 6pt after headings
- Use bullet points for achievements and responsibilities

STRUCTURE:
1. Header Section:
   - Full Name (centered, bold, 16pt)
   - Contact information below name (email, phone, LinkedIn, location)
   - Professional summary (2-3 sentences highlighting key strengths)

2. Professional Experience:
   - Job Title | Company Name | Dates
   - 3-5 bullet points per role highlighting achievements with metrics
   - Use action verbs (Led, Developed, Implemented, Achieved)

3. Education:
   - Degree, Major | University Name | Graduation Year
   - Relevant coursework or honors (if applicable)

4. Skills:
   - Technical Skills: List relevant technical competencies
   - Soft Skills: Communication, Leadership, Problem-solving, etc.

5. Certifications (if applicable):
   - Certification Name | Issuing Organization | Year

Make the content comprehensive with realistic examples. Include quantifiable achievements (e.g., "Increased sales by 25%", "Managed team of 8 developers").""",
        icon="briefcase"
    ),
    Template(
        id="invoice",
        title="Professional Invoice",
        description="Clean and detailed invoice template for business transactions",
        category="Business",
        prompt="""Create a Professional Invoice using the following specifications:

FORMATTING REQUIREMENTS:
- Use Arial font, 11pt for body text
- Company header with logo placeholder and contact details
- Clear table structure for line items
- Professional color scheme (navy blue #1E3A8A for headers)
- Bold headings and proper alignment

STRUCTURE:
1. Header Section:
   - "INVOICE" title (centered, bold, 20pt)
   - Company Name and full contact details (address, phone, email, website)
   - Invoice Number: INV-2024-001
   - Invoice Date: [Current Date]
   - Due Date: [30 days from invoice date]

2. Bill To Section:
   - Client Name
   - Client Company
   - Client Address
   - Client Contact Information

3. Invoice Details Table:
   - Columns: Item/Description | Quantity | Unit Price | Total
   - Include 5-7 sample line items with realistic services/products
   - Each item should have detailed description

4. Summary Section:
   - Subtotal
   - Tax (10% or applicable rate)
   - Discount (if any)
   - Total Amount Due (bold, highlighted)

5. Payment Information:
   - Payment Terms: Net 30
   - Accepted Payment Methods
   - Bank Account Details (if applicable)

6. Footer:
   - Thank you message
   - Terms and conditions summary
   - Late payment policy

Make it comprehensive and professional with realistic business examples.""",
        icon="receipt"
    ),
    Template(
        id="legal-contract",
        title="Legal Contract",
        description="Formal legal agreement template with standard clauses",
        category="Legal",
        prompt="""Create a Legal Contract using the following specifications:

FORMATTING REQUIREMENTS:
- Use Times New Roman font, 12pt for body text
- Professional, formal layout with proper legal formatting
- Numbered sections and subsections (1.1, 1.2, etc.)
- Bold headings for major sections
- Justified text alignment for body paragraphs
- 1.5 line spacing for readability

STRUCTURE:
1. Title and Header:
   - "SERVICE AGREEMENT" or "CONTRACT FOR SERVICES" (centered, bold, 14pt)
   - Date of Agreement
   - Parties Involved (Full legal names and addresses)

2. Recitals Section:
   - WHEREAS clauses explaining the background and purpose

3. Main Body - Key Sections:
   
   3.1 DEFINITIONS
   - Define key terms used throughout the contract
   
   3.2 SCOPE OF SERVICES
   - Detailed description of services to be provided
   - Deliverables and specifications
   
   3.3 TERM AND TERMINATION
   - Contract duration
   - Termination conditions and notice periods
   - Effect of termination
   
   3.4 COMPENSATION AND PAYMENT
   - Payment amount and schedule
   - Payment method
   - Late payment provisions
   
   3.5 CONFIDENTIALITY
   - Non-disclosure obligations
   - Exceptions to confidentiality
   
   3.6 INTELLECTUAL PROPERTY
   - Ownership of work product
   - License grants (if any)
   
   3.7 WARRANTIES AND REPRESENTATIONS
   - Mutual warranties
   - Disclaimer of other warranties
   
   3.8 LIMITATION OF LIABILITY
   - Liability caps
   - Exclusions
   
   3.9 INDEMNIFICATION
   - Indemnity obligations of each party
   
   3.10 GENERAL PROVISIONS
   - Governing law and jurisdiction
   - Entire agreement clause
   - Amendment procedures
   - Severability
   - Force majeure
   - Notices

4. Signature Block:
   - Signature lines for all parties
   - Date fields
   - Witness signature lines (if required)

Make this comprehensive with detailed, realistic legal language appropriate for a professional services agreement.""",
        icon="file-text"
    ),
    Template(
        id="business-proposal",
        title="Business Proposal",
        description="Comprehensive proposal for business projects and partnerships",
        category="Business",
        prompt="""Create a Comprehensive Business Proposal using the following specifications:

FORMATTING REQUIREMENTS:
- Use Calibri or Arial font, 11pt for body text, larger for headings
- Professional cover page with title and company branding
- Clear hierarchy: Title (16pt) > Heading 1 (14pt) > Heading 2 (12pt) > Body (11pt)
- Use tables for budget breakdowns and timelines
- Include bullet points for key benefits and deliverables
- Color accents in blue (#2563EB) for headings and important elements

STRUCTURE:

1. COVER PAGE
   - Proposal Title
   - Prepared For: [Client Name]
   - Prepared By: [Your Company]
   - Date
   - Confidentiality Notice

2. EXECUTIVE SUMMARY (1-2 pages)
   - Brief overview of the proposal
   - Key benefits and value proposition
   - Investment required and expected ROI
   - Call to action

3. COMPANY OVERVIEW
   - About your company
   - Mission and vision
   - Core competencies
   - Relevant experience and past successes
   - Team credentials

4. PROBLEM STATEMENT
   - Current challenges faced by the client
   - Market analysis and context
   - Impact of not addressing the problem
   - Opportunity for improvement

5. PROPOSED SOLUTION
   - Detailed description of your solution
   - How it addresses each problem point
   - Unique value proposition
   - Technology/methodology to be used
   - Innovation and competitive advantages

6. IMPLEMENTATION PLAN
   - Phase-by-phase breakdown
   - Detailed timeline with milestones (use table format)
   - Resource allocation
   - Key deliverables for each phase
   - Quality assurance measures

7. BUDGET AND PRICING
   - Itemized cost breakdown (use professional table)
   - Payment schedule
   - What's included vs. optional add-ons
   - ROI analysis and cost-benefit comparison

8. RISK ASSESSMENT AND MITIGATION
   - Potential risks identified
   - Mitigation strategies for each risk
   - Contingency plans

9. SUCCESS METRICS
   - KPIs to measure success
   - Reporting frequency and format
   - Evaluation criteria

10. TERMS AND CONDITIONS
    - Project scope boundaries
    - Assumptions and dependencies
    - Change management process
    - Warranty and support terms

11. CONCLUSION AND NEXT STEPS
    - Summary of key points
    - Clear call to action
    - Contact information
    - Proposal validity period

Make this extensive and professional with realistic business scenarios, detailed timelines, and comprehensive budget tables.""",
        icon="presentation"
    ),
    Template(
        id="meeting-minutes",
        title="Meeting Minutes",
        description="Structured template for recording meeting discussions and action items",
        category="Professional",
        prompt="""Create Professional Meeting Minutes using the following specifications:

FORMATTING REQUIREMENTS:
- Use Arial or Calibri font, 11pt for body text
- Clear section headers in bold, 12pt
- Use tables for attendees and action items
- Bullet points for discussion points
- Consistent spacing and professional layout

STRUCTURE:

1. HEADER SECTION
   - Meeting Title (bold, 14pt)
   - Date and Time
   - Location/Platform (e.g., Conference Room A / Zoom)
   - Meeting Duration

2. ATTENDEES
   - Present: (list all attendees with titles)
   - Absent: (list with reasons if known)
   - Guests/Special Attendees: (if any)

3. AGENDA ITEMS
   - List of topics to be covered
   - Time allocation for each item

4. MEETING OBJECTIVES
   - Primary goals of the meeting
   - Expected outcomes

5. DISCUSSION POINTS (Organized by Agenda Item)
   For each agenda item:
   - Topic heading
   - Key points discussed (bullet points)
   - Decisions made
   - Concerns raised
   - Data/metrics shared

6. ACTION ITEMS TABLE
   Create a comprehensive table with columns:
   - Action Item Description
   - Assigned To
   - Due Date
   - Priority (High/Medium/Low)
   - Status
   
   Include 8-10 realistic action items

7. DECISIONS MADE
   - List all formal decisions
   - Voting results (if applicable)
   - Rationale for key decisions

8. OPEN ISSUES
   - Unresolved matters
   - Items requiring further discussion
   - Parking lot items for future meetings

9. NEXT MEETING
   - Proposed date and time
   - Tentative agenda items
   - Preparation required

10. ADDITIONAL NOTES
    - Important announcements
    - Resource links or documents referenced
    - Follow-up communications needed

11. APPROVAL SECTION
    - Minutes prepared by: [Name]
    - Date prepared:
    - Reviewed and approved by: [Name]
    - Approval date:

Make this detailed and realistic with substantive discussion points and actionable items.""",
        icon="clipboard"
    ),
    Template(
        id="project-charter",
        title="Project Charter",
        description="Formal document to authorize a project and define objectives",
        category="Professional",
        prompt="""Create a Comprehensive Project Charter using the following specifications:

FORMATTING REQUIREMENTS:
- Use Calibri font, 11pt for body text
- Professional header with project name and logo placeholder
- Clear section headings in bold, 13pt
- Use tables for stakeholders, milestones, and budget
- Color coding: Blue (#2563EB) for headings, Green for success criteria
- Proper spacing: 12pt after paragraphs

STRUCTURE:

1. PROJECT OVERVIEW
   - Project Name (bold, 16pt)
   - Project Code/ID
   - Project Manager
   - Sponsor
   - Date Created
   - Version Number

2. EXECUTIVE SUMMARY
   - Brief project description (2-3 paragraphs)
   - Strategic alignment with organizational goals
   - Expected business value

3. PROJECT PURPOSE AND JUSTIFICATION
   - Business case
   - Problem or opportunity being addressed
   - Expected benefits and ROI
   - Alignment with strategic objectives

4. PROJECT OBJECTIVES (SMART Format)
   - Specific, measurable objectives (5-7 objectives)
   - Success criteria for each objective
   - Timeline for achievement

5. PROJECT SCOPE
   - In Scope:
     * Detailed list of what's included
     * Deliverables
     * Features and functionality
   - Out of Scope:
     * Explicitly excluded items
     * Future phase considerations

6. STAKEHOLDER ANALYSIS
   Create a table with:
   - Stakeholder Name/Group
   - Role
   - Interest Level
   - Influence Level
   - Communication Needs

7. HIGH-LEVEL REQUIREMENTS
   - Functional requirements (8-10 items)
   - Technical requirements
   - Business requirements
   - Compliance/regulatory requirements

8. PROJECT MILESTONES AND TIMELINE
   Create a table with:
   - Milestone Name
   - Description
   - Target Date
   - Dependencies
   - Deliverables
   
   Include 10-12 major milestones

9. BUDGET SUMMARY
   Create a detailed table:
   - Cost Category
   - Estimated Cost
   - Notes
   
   Categories: Personnel, Technology, Equipment, Training, Contingency, etc.
   - Total Project Budget
   - Funding Source

10. RESOURCE REQUIREMENTS
    - Human resources (roles and FTE)
    - Technology and tools
    - Facilities and equipment
    - External vendors/contractors

11. ASSUMPTIONS
    - List 6-8 key assumptions
    - Impact if assumptions prove false

12. CONSTRAINTS
    - Time constraints
    - Budget constraints
    - Resource constraints
    - Technical constraints
    - Regulatory constraints

13. RISKS AND MITIGATION
    Create a table:
    - Risk Description
    - Probability (High/Medium/Low)
    - Impact (High/Medium/Low)
    - Mitigation Strategy

14. SUCCESS CRITERIA
    - Quantifiable metrics
    - Quality standards
    - Acceptance criteria
    - Performance benchmarks

15. PROJECT GOVERNANCE
    - Steering committee members
    - Decision-making authority
    - Escalation process
    - Reporting structure and frequency

16. APPROVAL AND SIGN-OFF
    - Project Sponsor signature line
    - Project Manager signature line
    - Key Stakeholder signature lines
    - Date fields

Make this comprehensive and professional with realistic project details, detailed tables, and substantive content for each section.""",
        icon="folder"
    )
]


@app.get("/")
async def root():
    """Root endpoint with API information"""
    return {
        "message": "AI Word Processor API",
        "version": "1.0.0",
        "endpoints": {
            "GET /templates": "Get available document templates",
            "POST /generate": "Generate a DOCX file from a prompt",
            "POST /export": "Export editor HTML to a DOCX file",
            "POST /export-pdf": "Export editor HTML to a PDF file",
            "POST /ai/edit": "Apply an AI edit instruction to the current HTML document",
            "GET /download/{filename}": "Download a generated file"
        }
    }


@app.get("/templates")
async def get_templates():
    """
    Get all available document templates.
    
    Returns:
        List of templates with their metadata
    """
    return {
        "success": True,
        "templates": [template.model_dump() for template in TEMPLATES]
    }


@app.post("/export", response_model=DocumentResponse)
async def export_document(request: ExportRequest):
    try:
        file_id = str(uuid.uuid4())
        raw_name = (request.filename or f"document_{file_id}.docx").strip() or f"document_{file_id}.docx"
        safe_name = Path(raw_name).name
        if not safe_name.lower().endswith(".docx"):
            safe_name = f"{safe_name}.docx"

        output_path = GENERATED_DIR / f"export_{file_id}_{safe_name}"
        data = html_to_docx_bytes(request.html)
        output_path.write_bytes(data)

        return DocumentResponse(
            success=True,
            message="Document exported successfully",
            download_url=f"/download/{output_path.name}",
            filename=output_path.name,
        )
    except Exception as e:
        error_trace = traceback.format_exc()
        print(f"Error exporting document: {error_trace}")
        return DocumentResponse(
            success=False,
            message="Failed to export document",
            error=str(e),
        )


@app.post("/export-pdf", response_model=DocumentResponse)
async def export_document_as_pdf(request: ExportRequest):
    try:
        file_id = str(uuid.uuid4())
        raw_name = (request.filename or f"document_{file_id}.pdf").strip() or f"document_{file_id}.pdf"
        safe_name = Path(raw_name).name
        if not safe_name.lower().endswith(".pdf"):
            safe_name = f"{safe_name}.pdf"

        output_path = GENERATED_DIR / f"export_{file_id}_{safe_name}"
        data = html_to_pdf_bytes(request.html)
        output_path.write_bytes(data)

        return DocumentResponse(
            success=True,
            message="PDF exported successfully",
            download_url=f"/download/{output_path.name}",
            filename=output_path.name,
        )
    except Exception as e:
        error_trace = traceback.format_exc()
        print(f"Error exporting PDF: {error_trace}")
        return DocumentResponse(
            success=False,
            message="Failed to export PDF",
            error=str(e),
        )


@app.post("/ai/edit", response_model=AiEditResponse)
async def ai_edit_document(request: AiEditRequest):
    try:
        if not request.instruction.strip():
            return AiEditResponse(success=False, message="Missing instruction", error="Instruction is required")

        prompt = (
            "Update the following HTML document according to the instruction.\n\n"
            f"INSTRUCTION:\n{request.instruction}\n\n"
            f"HTML:\n{request.html}\n"
        )

        max_retries = 3
        last_error = None

        for attempt in range(max_retries):
            try:
                content_to_send = prompt
                if attempt > 0:
                    content_to_send = (
                        "The previous attempt failed. Fix the output and return ONLY valid HTML.\n\n"
                        f"ERROR:\n{str(last_error)}\n\n"
                        f"INSTRUCTION:\n{request.instruction}\n\n"
                        f"HTML:\n{request.html}\n"
                    )

                response = client.models.generate_content(
                    model=MODEL_NAME,
                    contents=content_to_send,
                    config=types.GenerateContentConfig(
                        system_instruction=EDIT_SYS_INSTRUCT
                    )
                )

                updated_html = _strip_code_fences(response.text)
                if not updated_html:
                    raise Exception("Empty AI response")
                if "<" not in updated_html or ">" not in updated_html:
                    raise Exception("AI response was not HTML")

                return AiEditResponse(
                    success=True,
                    message="AI edit applied successfully",
                    updated_html=updated_html,
                )
            except Exception as e:
                last_error = e
                if attempt == max_retries - 1:
                    raise

        return AiEditResponse(success=False, message="AI edit failed", error=str(last_error))

    except Exception as e:
        error_trace = traceback.format_exc()
        print(f"Error applying AI edit: {error_trace}")
        return AiEditResponse(success=False, message="AI edit failed", error=str(e))


@app.post("/generate", response_model=DocumentResponse)
async def generate_document(
    prompt: str = Form(...),
    file: UploadFile | None = File(None)
):
    """
    Generate a DOCX document based on the user's prompt and optional source file.
    
    Args:
        prompt: User's prompt describing what to generate
        file: Optional source file (PDF, TXT, or DOCX) for context
        
    Returns:
        DocumentResponse with download URL and metadata
    """
    try:
        # Validate file if provided
        if file:
            validate_upload_file(file)
        
        # Generate unique filename
        file_id = str(uuid.uuid4())
        filename = f"document_{file_id}.docx"
        output_path = GENERATED_DIR / filename
        
        # Process uploaded file if provided
        document_context = ""
        if file and file.filename:
            file_bytes = await file.read()
            
            # Validate file size
            if len(file_bytes) > MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=400,
                    detail=f"File size exceeds maximum allowed size of {MAX_FILE_SIZE // (1024*1024)}MB"
                )
            
            file_ext = Path(file.filename).suffix.lower()
            
            if file_ext == ".pdf":
                # For PDFs, we'll use Gemini's document understanding
                # We'll pass it directly to Gemini instead of extracting text
                document_context = "[PDF document provided for context]"
            elif file_ext == ".txt":
                document_context = extract_text_from_txt(file_bytes)
            elif file_ext == ".docx":
                document_context = extract_text_from_docx(file_bytes)
        
        # Create prompt for Gemini
        if document_context and document_context != "[PDF document provided for context]":
            gemini_prompt = f"""Generate Python code using python-docx to create a DETAILED, LONG, and PROFESSIONALLY FORMATTED document based on the following requirements:

{prompt}

SOURCE DOCUMENT CONTEXT:
{document_context}

Requirements:
- Make the content extensive and ellaborate.
- Use the source document context to inform the generated content.
- Use professional formatting (headings, spacing, fonts).
- Ensure the code handles the saving to 'output_path'.
"""
        else:
            gemini_prompt = f"""Generate Python code using python-docx to create a DETAILED, LONG, and PROFESSIONALLY FORMATTED document based on the following requirements:

{prompt}

Requirements:
- Make the content extensive and ellaborate.
- Use professional formatting (headings, spacing, fonts).
- Ensure the code handles the saving to 'output_path'.
"""
        
        # Retry loop for robust generation
        max_retries = 3
        last_error = None
        generated_code = ""
        
        for attempt in range(max_retries):
            try:
                content_to_send = gemini_prompt
                if attempt > 0:
                    print(f"Attempt {attempt + 1}: Retrying due to error: {last_error}")
                    # Refine prompt with error information
                    content_to_send = f"""The previous Python code you generated failed with the following error:
{str(last_error)}

The code that failed was:
{generated_code}

Please FIX the error and provide the corrected, complete executable Python code.
Ensure all imports are correct and the file is saved to 'output_path'.
output_path variable is available in the environment, do not define it, just use it.
"""
                
                # Prepare content parts for Gemini
                content_parts = []
                
                # If we have a PDF file, include it using Gemini's document understanding
                if file and file.filename and Path(file.filename).suffix.lower() == ".pdf":
                    # Reset file pointer
                    await file.seek(0)
                    pdf_bytes = await file.read()
                    
                    # Add PDF as inline data
                    content_parts.append(
                        types.Part.from_bytes(
                            data=pdf_bytes,
                            mime_type='application/pdf',
                        )
                    )
                
                # Add the text prompt
                content_parts.append(content_to_send)
                
                response = client.models.generate_content(
                    model=MODEL_NAME,
                    contents=content_parts,
                    config=types.GenerateContentConfig(
                        system_instruction=DOC_GEN_SYS_INSTRUCT
                    )
                )

                raw_code = response.text.strip()
                
                # enhanced markdown cleaning
                generated_code = raw_code
                if generated_code.startswith("```python"):
                    generated_code = generated_code.replace("```python", "", 1)
                elif generated_code.startswith("```"):
                    generated_code = generated_code.replace("```", "", 1)
                
                if generated_code.endswith("```"):
                    generated_code = generated_code[:-3]
                
                generated_code = generated_code.strip()
                
                # Prepare execution environment
                exec_globals = {
                    "output_path": str(output_path),
                    "__builtins__": __builtins__
                }
                
                # Execute the generated code
                exec(generated_code, exec_globals)
                
                # Verify the file was created
                if not output_path.exists():
                    raise Exception("Code executed without error, but document file was not created at 'output_path'.")
                
                # If we get here, it worked
                break
                
            except Exception as e:
                last_error = e
                # If this was the last attempt, we let the exception bubble up to the main try/except
                if attempt == max_retries - 1:
                    raise e
        
        # Generate download URL
        download_url = f"/download/{filename}"
        
        return DocumentResponse(
            success=True,
            message="Document generated successfully",
            download_url=download_url,
            filename=filename,
            generated_code=generated_code
        )
        
    except Exception as e:
        error_trace = traceback.format_exc()
        print(f"Error generating document: {error_trace}")
        
        return DocumentResponse(
            success=False,
            message="Failed to generate document",
            error=str(e),
            generated_code=generated_code if 'generated_code' in locals() else None
        )


@app.get("/download/{filename}")
async def download_file(filename: str):
    """
    Download a generated file (DOCX or PDF).
    
    Args:
        filename: Name of the file to download
        
    Returns:
        FileResponse with the file
    """
    file_path = GENERATED_DIR / filename
    
    if not file_path.exists():
        raise HTTPException(
            status_code=404, 
            detail="This document has been deleted as it exceeded the 10-minute retention period. Please generate a new document."
        )
    
    media_type = "application/pdf" if filename.lower().endswith(".pdf") else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    return FileResponse(
        path=str(file_path),
        filename=filename,
        media_type=media_type
    )


@app.delete("/cleanup")
async def cleanup_old_files(max_age_hours: int = 24):
    """
    Clean up old generated files.
    
    Args:
        max_age_hours: Maximum age of files to keep (default: 24 hours)
        
    Returns:
        Number of files deleted
    """
    deleted_count = 0
    current_time = datetime.now().timestamp()
    max_age_seconds = max_age_hours * 3600
    
    for file_path in GENERATED_DIR.glob("*"):
        if file_path.suffix.lower() in [".docx", ".pdf"]:
            file_age = current_time - file_path.stat().st_mtime
            if file_age > max_age_seconds:
                file_path.unlink()
                deleted_count += 1
    
    return {
        "message": f"Cleaned up {deleted_count} old files",
        "deleted_count": deleted_count
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
